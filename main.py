import os
import requests
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bidi.algorithm import get_display
import arabic_reshaper
import re
from flask import Flask, request  # Add this line to import Flask

# Rest of your code remains the same
# API Keys and Tokens
GEMINI_API_KEY = 'AIzaSyA0mYC_EoL3bVMhVRc0CF70uoeGZzVf59g'
TELEGRAM_TOKEN = '7392459074:AAG9sixtiU91cl_qv8sqrw363ilh5PZftYo'












async def webhook(request):
    """Handle incoming Telegram updates via webhook."""
    if request.method == "POST":
        await application.update_queue.put(Update.de_json(request.get_json(force=True), application.bot))
        return "OK"
    return "Hello, this is your Telegram bot webhook!"

# Initialize Flask app
app = Flask(__name__)

# Initialize Telegram application
application = Application.builder().token(TELEGRAM_TOKEN).build()

# Add handlers
application.add_handler(CommandHandler("start", start))
application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

@app.route("/", methods=["POST"])
async def index():
    """Handle the webhook request from Telegram."""
    return await webhook(request)

if __name__ == "__main__":
    # For local development, you can use polling
    # application.run_polling(allowed_updates=Update.ALL_TYPES)
    
    # For production with webhook
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))











# Load Arabic font for PDF
pdfmetrics.registerFont(TTFont('Arabic', 'Amiri-Regular.ttf'))
pdfmetrics.registerFont(TTFont('ArabicBold', 'Amiri Bold.ttf'))

async def generate_research(topic: str) -> str:
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={GEMINI_API_KEY}"
    headers = {'Content-Type': 'application/json'}
    prompt = f"اكتب بحثًا منظمًا وموجزًا من صفحة واحدة باللغة العربية عن {topic}، بما في ذلك المراجع في النهاية. يجب أن يكون المحتوى حوالي 600 كلمة."
    data = {
        "contents": [{
            "parts": [{
                "text": prompt
            }]
        }]
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        content = response.json()['candidates'][0]['content']['parts'][0]['text']
        return content.strip()
    except requests.RequestException as e:
        print(f"Error generating research: {e}")
        return "حدث خطأ أثناء إنشاء البحث. يرجى المحاولة مرة أخرى لاحقًا."

def clean_text(text: str) -> str:
    # Remove any HTML-like tags
    cleaned_text = re.sub(r'<[^>]+>', '', text)
    # Replace special characters used for emphasis with spaces
    cleaned_text = re.sub(r'[&^%$#@!_()*&\-]', ' ', cleaned_text)
    return cleaned_text

def reshape_arabic_text(text: str) -> str:
    reshaped_text = arabic_reshaper.reshape(text)
    return get_display(reshaped_text)

def create_pdf(research_text: str, filename: str, topic: str):
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=1.5*cm, leftMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    story = []
    
    # Add Arabic styles with right-to-left direction
    arabic_style = ParagraphStyle(name='Arabic', fontName='Arabic', fontSize=12, alignment=2, leading=20, rightIndent=0.5*cm, leftIndent=0.5*cm)
    title_style = ParagraphStyle(name='Title', fontName='ArabicBold', fontSize=18, alignment=1, spaceAfter=0.5*cm, textColor=colors.darkblue)
    
    # Add title
    story.append(Paragraph(reshape_arabic_text("بحث عن: " + topic), title_style))
    story.append(Spacer(1, 0.5 * cm))
    
    # Clean and format paragraphs in Arabic
    cleaned_text = clean_text(research_text)
    for paragraph in cleaned_text.split('\n\n'):
        story.append(Paragraph(reshape_arabic_text(paragraph), arabic_style))
        story.append(Spacer(1, 0.3 * cm))
    
    # Add footer
    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Arabic', 6)
        footer_text = reshape_arabic_text("تم برمجة بواسطة الطالب / محمد خضر")
        canvas.drawCentredString(A4[0]/2, 0.5 * cm, footer_text)
        canvas.restoreState()
    
    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)

def create_doc(research_text: str, filename: str, topic: str):
    doc = Document()
    
    # Set document settings for right-to-left text
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # Set document language to Arabic and text direction
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    
    rtl = OxmlElement('w:bidi')
    rtl.set(qn('w:val'), '1')
    doc.styles['Normal']._element.rPr.append(rtl)
    
    # Add title
    title = doc.add_heading(reshape_arabic_text(f"بحث عن: {topic}"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Dark blue color
    
    # Clean and add paragraphs
    cleaned_text = clean_text(research_text)
    for paragraph in cleaned_text.split('\n\n'):
        para = doc.add_paragraph(reshape_arabic_text(paragraph))
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(10)
    
    # Add footer
    section = doc.sections[-1]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = reshape_arabic_text("تم برمجة بواسطة الطالب / محمد خضر")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.size = Pt(6)
    
    doc.save(filename)

def reshape_arabic_text(text: str) -> str:
    reshaped_text = arabic_reshaper.reshape(text)
    return get_display(reshaped_text)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    welcome_message = "🚀مرحبًا بك في بوت البحث العلمي! 📚✨\n" \
                      "البوت من برمجة الطالب / محمد خضر (بالفرقة الاولي) للتواصل معه @FigoMK\n\n" \
                      "للبدء، ما عليك سوى إرسال موضوع ترغب في عمل بحث عنه. " \
                      "سيقوم بإنشاء ورقة موجزة من صفحة واحدة لك باستخدام تقنية الذكاء الاصطناعي عالية المستوى.\n\n" \
                      "مثال: 'المحاسبة المالية وتطبيقها في الواقع'"
    await update.message.reply_text(welcome_message)

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    topic = update.message.text
    
    generating_message = f"📝 جارٍ إنشاء بحث عن: '{topic}'...\nقد يستغرق هذا بضع لحظات، يرجى الانتظار."
    await update.message.reply_text(generating_message)
    
    try:
        research_text = await generate_research(topic)
        
        safe_topic = ''.join(e for e in topic if e.isalnum() or e.isspace())
        pdf_filename = f"{safe_topic}_بحث.pdf"
        doc_filename = f"{safe_topic}_بحث.docx"
        
        create_pdf(research_text, pdf_filename, topic)
        create_doc(research_text, doc_filename, topic)
        
        await update.message.reply_document(document=open(pdf_filename, 'rb'), filename=pdf_filename)
        await update.message.reply_document(document=open(doc_filename, 'rb'), filename=doc_filename)
        
        success_message = "🎉 تم إنشاء البحث بنجاح!\n\n" \
                          "لقد أرسلت لك نسختين من الورقة بصيغتي PDF وDOCX. " \
                          "😊 لا تتردد في طلب موضوع آخر في أي وقت."
        await update.message.reply_text(success_message)
        
        os.remove(pdf_filename)
        os.remove(doc_filename)
    except Exception as e:
        error_message = f"عذرًا، حدث خطأ أثناء إنشاء البحث. يرجى المحاولة مرة أخرى لاحقًا.\nالخطأ: {str(e)}"
        await update.message.reply_text(error_message)


def main():
    application = Application.builder().token(TELEGRAM_TOKEN).build()
    
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()
